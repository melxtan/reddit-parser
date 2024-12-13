import csv
import io
import json
import logging
import os
from typing import Any, Callable, Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st
from docx import Document
from reddit_analysis import RedditAnalyzer, analyze_reddit_data
from scrape_reddit import ScrapeReddit

logger = logging.getLogger(__name__)


def initialize_app() -> None:
    """Initialize the Streamlit application with required environment variables and state."""
    # Setup langfuse
    os.environ["LANGFUSE_PUBLIC_KEY"] = "pk-lf-457d1b99-6acb-490c-a50c-71916af2b291"
    os.environ["LANGFUSE_SECRET_KEY"] = "sk-lf-33ed664d-236c-4c0b-b6d3-7b066a012c0a"
    os.environ["LANGFUSE_HOST"] = "https://us.cloud.langfuse.com"  # 🇺🇸 US region

    logging.basicConfig(
        level=st.secrets.get("LOG_LEVEL", "INFO"),
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    )

    if "post_data" not in st.session_state:
        st.session_state.post_data = None
    if "analysis_results" not in st.session_state:
        st.session_state.analysis_results = {}
    if "aws_creds" not in st.session_state:
        st.session_state.aws_creds = None
    if "task_containers" not in st.session_state:
        st.session_state.task_containers = {}
    if "debug_info" not in st.session_state:
        st.session_state.debug_info = {}


def render_search_interface() -> Tuple[str, Optional[str], Optional[str], str, str, int, bool, str]:
    """Render the search interface and return user inputs."""
    search_type = st.radio(
        "Select search type:",
        ["Search Query", "Subreddit"],
        help="Choose whether to search across all of Reddit or scrape a specific subreddit",
    )

    search_query = None
    subreddit_name = None
    
    if search_type == "Search Query":
        search_query = st.text_input("Enter a search query:", key="search_query")
    else:
        subreddit_name = st.text_input(
            "Enter subreddit name (without r/):",
            key="subreddit_name",
            help="Enter the name of the subreddit you want to scrape (e.g., 'Python' for r/Python)",
        )

    col1, col2 = st.columns(2)

    with col1:
        search_options = (
            ["relevance", "hot", "top", "new", "comments"]
            if search_type == "Search Query"
            else ["hot", "new", "top"]
        )

        sort_options = {
            "relevance": "Relevance",
            "hot": "Hot",
            "top": "Top",
            "new": "New",
            "comments": "Comment count",
        }

        search_option = st.selectbox(
            "Select sort option:",
            search_options,
            format_func=lambda x: sort_options[x],
            key="search_option",
        )

    with col2:
        time_options = {
            "all": "All time",
            "year": "Past year",
            "month": "Past month",
            "week": "Past week",
            "day": "Today",
            "hour": "Past hour",
        }

        time_filter = st.selectbox(
            "Select time filter:",
            list(time_options.keys()),
            index=1,
            format_func=lambda x: time_options[x],
            key="time_filter",
        )

    max_posts = st.number_input(
        "Maximum number of posts to scrape (0 for no limit):",
        min_value=0,
        value=10,
        key="max_posts",
    )

    use_api = st.checkbox(
        "Use Reddit API (faster, but may hit rate limits)",
        value=True,
        key="use_api",
    )

    log_level = st.selectbox(
        "Select log level:",
        ["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"],
        index=1,
        key="log_level",
    )

    return search_type, search_query, subreddit_name, search_option, time_filter, max_posts, use_api, log_level


def scrape_reddit_data(
    search_type: str,
    search_query: Optional[str],
    subreddit_name: Optional[str],
    search_option: str,
    time_filter: str,
    max_posts: int,
    use_api: bool,
    log_level: str,
) -> Optional[List[Dict[str, Any]]]:
    """Scrape Reddit data based on user inputs."""
    try:
        scraper = ScrapeReddit(
            use_api=use_api,
            log_level=logging.getLevelName(log_level),
            client_id=st.secrets["REDDIT_CLIENT_ID"],
            client_secret=st.secrets["REDDIT_CLIENT_SECRET"],
            user_agent=st.secrets["REDDIT_USER_AGENT"],
        )

        st.info(f"Fetching posts. Max posts: {max_posts}")

        if search_type == "Search Query":
            post_urls = scraper.get_posts(
                search_query,
                time_filter=time_filter,
                search_option=search_option,
                limit=max_posts,
            )
        else:
            post_urls = scraper.get_subreddit_posts(
                subreddit_name,
                search_option=search_option,
                time_filter=time_filter,
                limit=max_posts,
            )

        st.info(f"Fetching post info for {len(post_urls)} posts")
        post_data = scraper.get_reddit_post_info(post_urls)

        scraper.destroy()
        return post_data

    except Exception as exc:
        st.error(f"An error occurred: {str(exc)}")
        logger.exception("An error occurred during scraping:")
        return None


def display_data_summary(
    post_data: List[Dict[str, Any]],
    search_query: Optional[str],
    search_option: str,
    time_filter: str,
) -> pd.DataFrame:
    """Display summary of scraped data and return DataFrame."""
    flattened_data = []
    for post in post_data:
        flattened_data.append({
            "type": "post",
            "post_id": post["id"],
            "title": post["title"],
            "body": post["body"],
            "author": post["author"],
            "score": post["score"],
            "created_at": post["created_at"],
            "num_comments": post["num_comments"],
            "subreddit": post["subreddit"],
        })

        for comment in post["comments"]:
            flattened_data.append({
                "type": "comment",
                "post_id": post["id"],
                "title": "",
                "body": comment["body"],
                "author": comment["author"],
                "score": comment["score"],
                "created_at": comment["created_at"],
                "num_comments": None,
                "subreddit": post["subreddit"],
            })

    df = pd.DataFrame(flattened_data)

    st.subheader("Summary")
    search_info = (
        f"Search query: {search_query}"
        if search_query
        else f"Subreddit: r/{post_data[0]['subreddit']}"
    )
    st.write(search_info)
    st.write(f"Sort: {search_option}, Time filter: {time_filter}")

    posts_count = len(df[df["type"] == "post"])
    comments_count = len(df[df["type"] == "comment"])
    st.write(f"Number of posts: {posts_count}")
    st.write(f"Number of comments: {comments_count}")
    st.write(f"Average post score: {df[df['type'] == 'post']['score'].mean():.2f}")
    st.write(f"Average comment score: {df[df['type'] == 'comment']['score'].mean():.2f}")

    # Create preview with limited comments
    st.subheader("Preview (Posts with up to 2 comments)")
    preview_rows = create_preview_rows(post_data)
    preview_df = pd.DataFrame(preview_rows)
    st.dataframe(preview_df, hide_index=True)

    return df


def create_preview_rows(post_data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Create preview rows for display."""
    preview_rows = []
    for post in post_data:
        preview_rows.append({
            "type": "post",
            "title": post["title"],
            "body": truncate_text(post["body"], 100),
            "author": post["author"],
            "score": post["score"],
            "created_at": post["created_at"],
            "subreddit": post["subreddit"],
        })

        for comment in post["comments"][:2]:
            preview_rows.append({
                "type": "comment",
                "title": "",
                "body": truncate_text(comment["body"], 100),
                "author": comment["author"],
                "score": comment["score"],
                "created_at": comment["created_at"],
                "subreddit": post["subreddit"],
            })

    return preview_rows


def truncate_text(text: str, max_length: int) -> str:
    """Truncate text to specified length and add ellipsis if needed."""
    return f"{text[:max_length]}..." if len(text) > max_length else text


def create_download_buttons(
    df: pd.DataFrame,
    post_data: List[Dict[str, Any]],
    search_query: Optional[str],
    search_option: str,
    time_filter: str,
) -> str:
    """Create download buttons for data and return filename base."""
    col1, col2 = st.columns(2)

    subreddit_name = post_data[0]["subreddit"] if not search_query else None
    filename = generate_filename(search_query, subreddit_name, search_option, time_filter)

    with col1:
        json_str = json.dumps(post_data, indent=2, ensure_ascii=False)
        st.download_button(
            label="Download JSON",
            data=json_str,
            file_name=f"{filename}.json",
            mime="application/json",
            key="post_data_json",
        )

    with col2:
        csv_buffer = io.StringIO()
        df.to_csv(
            csv_buffer,
            index=False,
            encoding="utf-8-sig",
            sep=",",
            quoting=csv.QUOTE_ALL,
            escapechar="\\",
            doublequote=True,
        )
        csv_str = csv_buffer.getvalue()
        st.download_button(
            label="Download CSV",
            data=csv_str.encode("utf-8-sig"),
            file_name=f"{filename}.csv",
            mime="text/csv",
            key="post_data_csv",
        )

    return filename


def generate_filename(
    search_query: Optional[str],
    subreddit_name: Optional[str],
    search_option: str,
    time_filter: str,
) -> str:
    """Generate a filename for downloaded data."""
    if search_query:
        safe_query = search_query.replace(" ", "_").lower()[:30]
        return f"reddit_{safe_query}_{search_option}_{time_filter}"
    return f"reddit_r_{subreddit_name}_{search_option}_{time_filter}"


def handle_aws_credentials() -> None:
    """Handle AWS credentials input and storage."""
    if not st.session_state.aws_creds:
        with st.form("aws_creds_form"):
            aws_access_key = st.text_input(
                "AWS Access Key ID",
                type="password",
                key="aws_access_key",
            )
            aws_secret_key = st.text_input(
                "AWS Secret Access Key",
                type="password",
                key="aws_secret_key",
            )
            aws_region = st.text_input("AWS Region", value="us-west-2", key="aws_region")

            if st.form_submit_button("Save AWS Credentials"):
                if aws_access_key and aws_secret_key:
                    st.session_state.aws_creds = {
                        "access_key": aws_access_key,
                        "secret_key": aws_secret_key,
                        "region": aws_region,
                    }
                    st.success("AWS credentials saved!")
                else:
                    st.error("Please enter both AWS Access Key ID and Secret Access Key")
    else:
        st.success("AWS credentials are set")
        if st.button("Clear AWS Credentials", key="clear_creds"):
            st.session_state.aws_creds = None
            st.rerun()


def create_task_containers(task_order: List[str]) -> None:
    """Create containers for each analysis task."""
    for task_name in task_order:
        task_title = task_name.replace("_", " ").title()

        st.subheader(task_title)
        status_container = st.empty()
        status_container.info(f"Running {task_title}...")
        result_container = st.empty()
        debug_container = st.expander("Show Request Details")
        st.write("---")
        st.session_state.task_containers[task_name] = {
            "status": status_container,
            "result": result_container,
            "debug": debug_container,
        }


def clean_xml_result(result: Dict[str, Any]) -> str:
    """Clean and format the analysis result for display."""
    if not isinstance(result, dict) or "analysis" not in result:
        return "Invalid result format"

    analysis = result["analysis"]
    if not isinstance(analysis, str):
        analysis = str(analysis)

    cleaned = analysis.strip()
    if (
        cleaned.startswith("<")
        and cleaned.endswith(">")
        and ">" in cleaned[1:]
        and "</" in cleaned
    ):
        first_close = cleaned.find(">")
        last_open = cleaned.rfind("</")
        if first_close != -1 and last_open != -1:
            cleaned = cleaned[first_close + 1 : last_open].strip()

    return cleaned

def update_task_status(
    task_name: str,
    result: Dict[str, Any],
    task_order: List[str],
    filename: str,
) -> None:
    """Update the status and display of a task."""
    containers = st.session_state.task_containers[task_name]

    if "error" in result:
        containers["status"].error(
            f"Error in {task_name.replace('_', ' ').title()}: {result['error']}"
        )
        return

    containers["status"].success(f"{task_name.replace('_', ' ').title()} completed!")
    st.session_state.analysis_results[task_name] = result
    cleaned_analysis = clean_xml_result(result)
    containers["result"].text_area(
        label="Analysis Result",
        value=cleaned_analysis,
        height=300,
        disabled=True,
        label_visibility="collapsed",
        key=f"text_area_{task_name}",
    )

    if "request_body" in result:
        with containers["debug"]:
            st.subheader("Request Details")
            st.json(result["request_body"])

    if task_name == task_order[-1]:
        st.rerun()


def run_analysis(
    post_data: List[Dict[str, Any]],
    search_query: str,
    task_order: List[str],
    filename: str,
    min_comment_score: int,
    num_top_posts: int,
) -> None:
    """Run analysis on Reddit post data."""
    try:
        os.environ["AWS_ACCESS_KEY_ID"] = st.session_state.aws_creds["access_key"]
        os.environ["AWS_SECRET_ACCESS_KEY"] = st.session_state.aws_creds["secret_key"]

        st.info(
            f"Due to rate limit, we are currently only analyzing top {num_top_posts} posts "
            f"with highest scores. Only comments with a minimum score of {min_comment_score} "
            "will be included in the analysis."
        )

        create_task_containers(task_order)
        st.session_state.analysis_results = {}
        st.session_state.debug_info = {}

        def callback(task_name: str, result: Dict[str, Any]) -> None:
            update_task_status(task_name, result, task_order, filename)

        analyze_reddit_data(
            post_data=post_data,
            search_query=search_query,
            callback=callback,
            region_name=st.session_state.aws_creds["region"],
            rate_limit_per_second=0.5,
            num_top_posts=num_top_posts,
            min_comment_score=min_comment_score,
        )

    except Exception as exc:
        st.error(f"Analysis failed: {str(exc)}")
        logger.exception("Analysis error:")


def create_word_doc(task_order: List[str]) -> io.BytesIO:
    """Create a Word document from analysis results."""
    doc = Document()
    for task_name in task_order:
        if task_name in st.session_state.analysis_results:
            result = st.session_state.analysis_results[task_name]
            if "error" not in result:
                doc.add_heading(task_name.replace("_", " ").title(), level=1)
                doc.add_paragraph(clean_xml_result(result))
                doc.add_paragraph()

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


def display_analysis_results(task_order: List[str], filename: str) -> None:
    """Display analysis results and create download buttons."""
    for task_name in task_order:
        if task_name not in st.session_state.analysis_results:
            continue

        result = st.session_state.analysis_results[task_name]
        if "error" in result:
            continue

        st.subheader(task_name.replace("_", " ").title())
        st.text_area(
            label="Analysis Result",
            value=clean_xml_result(result),
            height=300,
            label_visibility="collapsed",
            key=f"text_area_{task_name}",
        )

        if "request_body" in result:
            with st.expander("Show Request Details"):
                st.subheader("Request Details")
                st.json(result["request_body"])

        st.write("---")

    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            label="Download Complete Analysis (JSON)",
            data=json.dumps(st.session_state.analysis_results, indent=2, ensure_ascii=False),
            file_name=f"{filename}_analysis.json",
            mime="application/json",
            key="analysis_json_final",
        )

    with col2:
        st.download_button(
            label="Download Analysis (Word)",
            data=create_word_doc(task_order),
            file_name=f"{filename}_analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="analysis_word_final",
        )

    if st.button("Run New Analysis", key="run_new_analysis"):
        st.session_state.post_data = None
        st.session_state.analysis_results = {}
        st.session_state.task_containers = {}
        st.session_state.debug_info = {}
        st.rerun()


def main() -> None:
    """Main application function."""
    task_order = RedditAnalyzer.TASKS
    initialize_app()

    with st.sidebar:
        password_input = st.text_input(
            "Enter password to access the app:",
            type="password",
            key="password_input",
        )

        if password_input == st.secrets["APP_PASSWORD"]:
            st.subheader("AWS Credentials")
            handle_aws_credentials()

    if password_input != st.secrets["APP_PASSWORD"]:
        st.error("Incorrect password. Access denied.")
        return

    st.title("Reddit Post Scraper")
    (
        search_type,
        search_query,
        subreddit_name,
        search_option,
        time_filter,
        max_posts,
        use_api,
        log_level,
    ) = render_search_interface()

    if not st.button("Scrape", key="scrape_button"):
        if st.session_state.post_data:
            display_existing_data(
                search_type,
                search_query,
                subreddit_name,
                search_option,
                time_filter,
                task_order,
            )
        return

    # Validate input based on search type
    if not validate_search_input(search_type, search_query, subreddit_name):
        st.warning(
            "Please enter a search query or subreddit name, depending on your selected search type."
        )
        return

    with st.spinner("Scraping data..."):
        post_data = scrape_reddit_data(
            search_type,
            search_query,
            subreddit_name,
            search_option,
            time_filter,
            max_posts,
            use_api,
            log_level,
        )
        if post_data:
            st.session_state.post_data = post_data
            st.session_state.analysis_results = {}
            st.session_state.task_containers = {}
            st.session_state.debug_info = {}


def validate_search_input(
    search_type: str,
    search_query: Optional[str],
    subreddit_name: Optional[str],
) -> bool:
    """Validate search input based on search type."""
    return (
        (search_type == "Search Query" and search_query)
        or (search_type == "Subreddit" and subreddit_name)
    )


def display_existing_data(
    search_type: str,
    search_query: Optional[str],
    subreddit_name: Optional[str],
    search_option: str,
    time_filter: str,
    task_order: List[str],
) -> None:
    """Display existing data and analysis options."""
    df = display_data_summary(
        st.session_state.post_data,
        search_query if search_type == "Search Query" else None,
        search_option,
        time_filter,
    )
    filename = create_download_buttons(
        df,
        st.session_state.post_data,
        search_query if search_type == "Search Query" else None,
        search_option,
        time_filter,
    )

    if not st.session_state.aws_creds:
        return

    st.subheader("Reddit Post Analysis")
    analysis_params = get_analysis_parameters()
    
    if st.button("Analyze Reddit Posts"):
        run_analysis(
            st.session_state.post_data,
            search_type,
            search_query,
            subreddit_name,
            task_order,
            filename,
            min_comment_score=analysis_params["min_comment_score"],
            num_top_posts=analysis_params["num_top_posts"],
        )
    elif st.session_state.analysis_results:
        display_analysis_results(task_order, filename)


def get_analysis_parameters() -> Dict[str, int]:
    """Get analysis parameters from user input."""
    col1, col2 = st.columns(2)
    
    with col1:
        min_comment_score = st.number_input(
            "Minimum comment score to include in analysis:",
            min_value=-100,
            value=1,
            help="Comments on Reddit start with a score of 1 (author's automatic upvote). "
            "A score of 0 means one downvote, negative scores mean more downvotes than upvotes. "
            "Only comments with scores >= this value will be analyzed.",
            key="min_comment_score",
        )
    
    with col2:
        num_posts = len(st.session_state.post_data)
        num_top_posts = st.number_input(
            "Number of top posts to analyze:",
            min_value=1,
            max_value=num_posts,
            value=min(10, num_posts),
            help="Select how many of the top posts (sorted by score) to include in the analysis.",
            key="num_top_posts",
        )

    return {
        "min_comment_score": min_comment_score,
        "num_top_posts": num_top_posts,
    }


if __name__ == "__main__":
    main()
